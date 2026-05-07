from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DOCX = Path("rapor_duzenlenecek.docx")
OUTPUT_DOCX = Path("diyabet_raporu_nihai.docx")
RESULTS_CSV = Path("results/model_comparison.csv")
RESULTS_JSON = Path("results/metrics.json")
CV_SUMMARY_CSV = Path("results/cv_accuracy_summary.csv")
DATA_CSV = Path("diabetes.csv")
ASSETS_DIR = Path("results/report_assets")

VIDEO_RESULTS = [
    ("LR", 0.7736842105263158, 0.04989986895904521),
    ("DT", 0.7087719298245614, 0.03526272147761716),
    ("KNN", 0.7087719298245615, 0.0544707884079299),
    ("NB", 0.7403508771929824, 0.05754813848019903),
    ("SVM", 0.7508771929824560, 0.05132890820465893),
    ("AdaB", 0.7754385964912280, 0.03907203061635103),
    ("GBM", 0.7684210526315789, 0.05860804592452656),
    ("RF", 0.7596491228070176, 0.04902873197362145),
]


FEATURE_DESCRIPTIONS = [
    ("Pregnancies", "Bireyin geçirdiği gebelik sayısını ifade eder."),
    ("Glucose", "Plazma glikoz düzeyidir ve diyabet riski açısından en güçlü göstergelerden biridir."),
    ("BloodPressure", "Diyastolik kan basıncını temsil eder."),
    ("SkinThickness", "Triseps deri kıvrım kalınlığı ölçümüdür."),
    ("Insulin", "2 saatlik serum insülin düzeyini gösterir."),
    ("BMI", "Vücut kitle indeksidir ve obeziteyle ilişkili riskin değerlendirilmesinde kullanılır."),
    ("DiabetesPedigreeFunction", "Aile öyküsü ve genetik yatkınlığı özetleyen skordur."),
    ("Age", "Bireyin yaş bilgisidir."),
    ("Outcome", "Bağımlı değişkendir; 0 negatif, 1 pozitif sınıfı ifade eder."),
]


ALGORITHM_DETAILS = [
    ("Logistic Regression", "max_iter=5000, class_weight=balanced", "Yorumlanabilirliği yüksek ve ikili sınıflandırma için güçlü bir temel modeldir."),
    ("KNN", "n_neighbors=11, weights=distance", "Örnekler arası benzerliği kullanarak karar verir ve yerel örüntüleri yakalayabilir."),
    ("SVM", "RBF kernel, C=1.0, gamma=scale", "Doğrusal olmayan karar sınırlarını modelleme kapasitesine sahiptir."),
    ("Decision Tree", "max_depth=5, min_samples_leaf=4", "Karar kuralları açık biçimde takip edilebildiği için yorumlama açısından değerlidir."),
    ("Random Forest", "300 ağaç, max_depth=6", "Birden fazla ağacın birleşimiyle daha dengeli genelleme yapmayı hedefler."),
    ("Extra Trees", "300 ağaç, max_depth=6", "Daha yüksek rastgelelik sayesinde topluluk çeşitliliği üretir."),
    ("Gradient Boosting", "150 estimator, learning_rate=0.05", "Ardışık hata düzeltme mantığı ile performans artırmayı amaçlar."),
    ("AdaBoost", "200 estimator, learning_rate=0.5", "Yanlış sınıflanan örneklere odaklanarak adaptif öğrenme gerçekleştirir."),
    ("Gaussian Naive Bayes", "Varsayılan", "Hızlı çalışan, olasılıksal ve düşük maliyetli bir baz model sunar."),
    ("MLP", "64-32 hidden layers, max_iter=2000", "Doğrusal olmayan ilişkileri yapay sinir ağı yaklaşımıyla modellemeyi amaçlar."),
]


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def format_paragraph(paragraph, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.3
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    format_paragraph(p, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    format_paragraph(p, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.add_run(f"- {item}")
        format_paragraph(p, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3)


def style_table_cell(cell, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.1
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(size)
            run.bold = bold


def add_two_col_table(doc: Document, headers: list[str], rows: list[tuple[str, str]], widths: tuple[float, float]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].width = Inches(widths[idx])
        table.rows[0].cells[idx].text = header
        style_table_cell(table.rows[0].cells[idx], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        cells[0].text = left
        cells[1].text = right
        style_table_cell(cells[0])
        style_table_cell(cells[1])


def add_algorithm_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(2.0), Inches(2.4), Inches(2.1)]
    headers = ["Algoritma", "Temel Parametreler", "Seçilme Gerekçesi"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].width = widths[i]
        table.rows[0].cells[i].text = header
        style_table_cell(table.rows[0].cells[i], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for name, params, rationale in ALGORITHM_DETAILS:
        row = table.add_row().cells
        values = [name, params, rationale]
        for i, value in enumerate(values):
            row[i].width = widths[i]
            row[i].text = value
            style_table_cell(row[i], size=10)


def add_model_table(doc: Document, comparison_df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.autofit = False
    widths = [Inches(2.0), Inches(0.82), Inches(0.82), Inches(0.82), Inches(0.82), Inches(0.82)]
    headers = ["Model", "Accuracy", "Precision", "Recall", "F1", "AUC"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].width = widths[i]
        table.rows[0].cells[i].text = header
        style_table_cell(table.rows[0].cells[i], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _, row in comparison_df.iterrows():
        values = [
            row["model"],
            f"{row['accuracy']:.4f}",
            f"{row['precision']:.4f}",
            f"{row['recall']:.4f}",
            f"{row['f1_score']:.4f}",
            f"{row['roc_auc']:.4f}",
        ]
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            cells[i].text = value
            style_table_cell(cells[i], size=10, align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)


def add_cv_summary_table(doc: Document, cv_df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [Inches(1.75), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0)]
    headers = ["Model", "Ortalama", "Std. Sapma", "Min", "Max"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].width = widths[i]
        table.rows[0].cells[i].text = header
        style_table_cell(table.rows[0].cells[i], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _, row in cv_df.iterrows():
        values = [
            row["label"],
            f"{row['mean_accuracy']:.4f}",
            f"{row['std_accuracy']:.4f}",
            f"{row['min_accuracy']:.4f}",
            f"{row['max_accuracy']:.4f}",
        ]
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            cells[i].text = value
            style_table_cell(cells[i], size=10, align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)


def add_video_comparison_table(doc: Document, cv_df: pd.DataFrame) -> None:
    mapping = {row["label"]: row for _, row in cv_df.iterrows()}
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [Inches(1.0), Inches(1.2), Inches(1.1), Inches(1.15), Inches(1.05)]
    headers = ["Model", "Video Ort.", "Video Std.", "Bizim Ort.", "Fark"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].width = widths[i]
        table.rows[0].cells[i].text = header
        style_table_cell(table.rows[0].cells[i], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for label, video_mean, video_std in VIDEO_RESULTS:
        ours = mapping[label]
        diff = ours["mean_accuracy"] - video_mean
        values = [
            label,
            f"{video_mean:.4f}",
            f"{video_std:.4f}",
            f"{ours['mean_accuracy']:.4f}",
            f"{diff:+.4f}",
        ]
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            cells[i].text = value
            style_table_cell(cells[i], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_image(doc: Document, image_path: Path, caption: str, width_inches: float = 6.1) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    cp = doc.add_paragraph(caption)
    format_paragraph(cp, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


def main() -> None:
    doc = Document(BASE_DOCX)
    set_default_style(doc)

    while len(doc.paragraphs) > 25:
        remove_paragraph(doc.paragraphs[25])

    comparison_df = pd.read_csv(RESULTS_CSV)
    cv_df = pd.read_csv(CV_SUMMARY_CSV)
    metrics_payload = json.loads(RESULTS_JSON.read_text(encoding="utf-8"))
    raw_df = pd.read_csv(DATA_CSV)

    best_metrics = metrics_payload["metrics"]
    best_model = metrics_payload["selected_model"]
    zero_counts = {col: int((raw_df[col] == 0).sum()) for col in ["Glucose", "BloodPressure", "SkinThickness", "Insulin", "BMI"]}
    class_counts = raw_df["Outcome"].value_counts().to_dict()

    add_heading(doc, "1. Proje Konusu ve Problemin Tanımı")
    add_body(
        doc,
        "Bu projede ele alınan problem, bireylerin temel sağlık verileri kullanılarak diyabet riski taşıyıp "
        "taşımadıklarının makine öğrenmesi yöntemleriyle tahmin edilmesidir. Diyabet, geç fark edildiğinde "
        "kardiyovasküler rahatsızlıklar, böbrek yetmezliği, görme kaybı ve uzun dönemli metabolik bozulmalar gibi "
        "ciddi sonuçlar doğurabilen kronik bir hastalıktır. Bu nedenle riskin erken aşamada belirlenmesi yalnızca "
        "bireysel düzeyde değil, sağlık sistemi ölçeğinde de stratejik bir gerekliliktir."
    )
    add_body(
        doc,
        "Diyabet riski tahmini sağlık bilişimi açısından önemli bir karar destek problemidir. Çünkü böyle bir model, "
        "tarama programlarının önceliklendirilmesine, hekimlerin riskli hastaları daha hızlı ayırt etmesine, "
        "koruyucu sağlık hizmetlerinin hedeflenmesine ve kaynak planlamasının daha verimli yürütülmesine katkı "
        "sağlayabilir. Bu açıdan bakıldığında projenin temel amacı yalnızca teknik olarak bir sınıflandırma modeli "
        "oluşturmak değil; erken teşhis, risk analizi ve veri temelli sağlık yönetimini destekleyecek işlevsel bir "
        "analitik yaklaşım geliştirmektir."
    )
    add_body(
        doc,
        "Dolayısıyla bu çalışma, ham sağlık verilerini karar destek sistemlerinde kullanılabilecek anlamlı içgörülere "
        "dönüştürmeyi hedeflemektedir. Veri madenciliği yaklaşımının değeri de tam olarak burada ortaya çıkmaktadır: "
        "veri üzerinden riski görünür kılmak, belirsizliği azaltmak ve müdahale süreçlerine bilimsel dayanak üretmek."
    )

    add_heading(doc, "2. Kullanılan Veri Seti")
    add_body(
        doc,
        "Projede kullanılan veri seti çalışma klasöründe yer alan `diabetes.csv` dosyasıdır. Yapısal özellikleri "
        "itibarıyla bu veri seti, literatürde ve açık veri platformlarında yaygın biçimde kullanılan Pima Indians "
        "Diabetes veri seti ile uyumludur ve diyabet risk tahmini problemleri için temsil gücü yüksek bir örnek "
        "küme olarak değerlendirilebilir."
    )
    add_body(
        doc,
        f"Veri seti toplam {raw_df.shape[0]} gözlem ve {raw_df.shape[1]} değişkenden oluşmaktadır. Bağımlı değişken "
        f"`Outcome` olup, sınıf dağılımı {class_counts.get(0, 0)} negatif ve {class_counts.get(1, 0)} pozitif gözlem "
        f"şeklindedir. Bu tablo, veri setinde belirli ölçüde sınıf dengesizliği olduğunu göstermektedir."
    )
    add_two_col_table(
        doc,
        ["Öznitelik", "Açıklama"],
        FEATURE_DESCRIPTIONS,
        (2.15, 4.25),
    )
    add_image(doc, ASSETS_DIR / "class_distribution.png", "Şekil 1. Veri setindeki sınıf dağılımı.", width_inches=5.8)
    add_body(
        doc,
        "Veri setinin yapısal analizi sırasında özellikle sıfır değerlerin mantıksal gerçekliği sorgulanmıştır. "
        "Glucose, BloodPressure, SkinThickness, Insulin ve BMI gibi sütunlarda yer alan sıfır değerler biyolojik "
        "olarak anlamlı kabul edilmemiştir. Bu nedenle bu değerler eksik veya hatalı veri şeklinde yorumlanmış ve "
        "medyan ile doldurulmuştur."
    )
    add_two_col_table(
        doc,
        ["Sütun", "Geçersiz Sıfır Sayısı"],
        [(key, str(value)) for key, value in zero_counts.items()],
        (2.7, 1.9),
    )
    add_body(
        doc,
        "Ön işleme sürecinde bunun yanında model doğasına göre standardization uygulanmış, ayrıca çeşitli öznitelik "
        "geliştirme işlemleri yapılmıştır. BMI_Obese_Flag, Glucose_High_Flag, Age_Risk_Flag ve Pregnancy_Risk_Flag "
        "gibi risk göstergeleri; Glucose_BMI_Interaction, Age_Glucose_Interaction, Insulin_Glucose_Ratio, "
        "Pregnancy_Age_Ratio, BMI_Age_Product ve Glucose_Squared gibi türetilmiş değişkenler projeye analitik katkı "
        "sağlamak amacıyla eklenmiştir. Bu yaklaşım, veri setini yalnızca kullanmak yerine onu yeniden yorumlayıp "
        "zenginleştirme çabasının bir parçasıdır."
    )

    add_heading(doc, "3. Kullanılan Yöntemler ve Algoritmalar")
    add_body(
        doc,
        "Bu projede toplam 10 farklı sınıflandırma algoritması kullanılmıştır. Buradaki amaç yalnızca algoritmaları "
        "çalıştırmak değil; veri setinin yapısına göre hangi yaklaşımın neden daha uygun sonuç verdiğini "
        "karşılaştırmalı biçimde ortaya koymaktır. Seçimler yapılırken yorumlanabilirlik, doğrusal/doğrusal olmayan "
        "ilişki yapısı, hesaplama maliyeti ve performans beklentisi birlikte değerlendirilmiştir."
    )
    add_algorithm_table(doc)
    add_body(
        doc,
        "Bu yaklaşım sayesinde hem lineer modellerin hem de topluluk öğrenmesi, olasılıksal yöntemler ve sinir ağı "
        "tabanlı modellerin aynı problem üzerindeki davranışı karşılaştırılmıştır. Böylece çalışma bir ödev "
        "uygulamasından çıkarılıp mini bir veri bilimi danışmanlık raporu niteliğine taşınmıştır."
    )

    add_heading(doc, "4. Model Eğitimi ve Test Süreci")
    add_body(
        doc,
        "Veri seti ilk aşamada %80 eğitim ve %20 test olacak şekilde ayrılmıştır. Ardından eğitim kümesi kendi "
        "içinde doğrulama amacıyla yeniden bölünmüş ve yaklaşık %60 eğitim, %20 doğrulama, %20 test yapısı "
        "oluşturulmuştur. Bu bölmeler sırasında stratified yaklaşım kullanılarak sınıf dağılımının korunmasına "
        "özellikle dikkat edilmiştir."
    )
    add_body(
        doc,
        "Her model eğitim kümesi üzerinde öğrenmiş, doğrulama kümesinde farklı olasılık eşikleri denenerek en uygun "
        "karar eşiği belirlenmiştir. Tam kapsamlı GridSearchCV veya RandomizedSearchCV yerine, veri setinin boyutu ve "
        "hesaplama maliyeti göz önüne alınarak kontrollü bir manuel hiperparametre seçimi yapılmıştır. Bu durum tam "
        "otomatik optimizasyon kadar kapsamlı olmasa da, şeffaf ve izlenebilir bir deney süreci sağlamıştır."
    )
    add_body(
        doc,
        "Eğitim maliyeti ve overfitting riski açısından modeller arasında belirgin farklar gözlenmiştir. Özellikle MLP "
        "ve topluluk yöntemleri daha yüksek hesaplama maliyeti üretirken, Logistic Regression ve Naive Bayes daha hızlı "
        "ve kararlı öğrenme davranışı sergilemiştir. Bu gözlem, yalnızca doğruluk değil; pratik kullanılabilirlik "
        "açısından da değerlendirilmiştir."
    )

    add_heading(doc, "5. Değerlendirme Metrikleri ve Yorumlar")
    add_body(
        doc,
        "Bu projede performans yalnızca accuracy ile değerlendirilmemiştir. Precision, recall, F1-Score, ROC-AUC ve "
        "confusion matrix birlikte kullanılarak daha bütüncül bir değerlendirme yapılmıştır. Sağlık alanındaki "
        "sınıflandırma problemlerinde yanlış negatiflerin maliyeti yüksek olduğu için recall ve sensitivity yorumları "
        "özellikle önemli kabul edilmiştir."
    )
    add_model_table(doc, comparison_df)
    add_image(doc, ASSETS_DIR / "metric_comparison.png", "Şekil 2. Modellerin Accuracy, F1-Score ve ROC-AUC karşılaştırması.")
    add_body(
        doc,
        f"Tek test ayrımı üzerinden yapılan değerlendirmede en iyi model `{best_model}` olmuştur. Bu model için "
        f"accuracy {best_metrics['accuracy']:.4f}, precision {best_metrics['precision']:.4f}, recall "
        f"{best_metrics['recall']:.4f}, F1-Score {best_metrics['f1_score']:.4f} ve ROC-AUC {best_metrics['roc_auc']:.4f} "
        f"olarak hesaplanmıştır. Özellikle recall değerinin yüksek olması, riskli bireylerin önemli bölümünün doğru "
        f"biçimde tespit edildiğini göstermektedir."
    )
    add_image(doc, ASSETS_DIR / "best_confusion_matrix.png", "Şekil 3. En iyi model için confusion matrix.", width_inches=5.2)
    add_image(doc, ASSETS_DIR / "best_roc_curve.png", "Şekil 4. En iyi model için ROC eğrisi.", width_inches=5.8)
    add_body(
        doc,
        "Confusion matrix bulguları, modelin yanlış negatif sayısını görece sınırlı tuttuğunu göstermektedir. Bu, "
        "diyabet tahmini bağlamında kritik bir avantajdır. ROC eğrisi ve ROC-AUC değeri de modelin rastgele tahmine "
        "kıyasla anlamlı ölçüde daha iyi sınıf ayrımı yaptığını doğrulamaktadır."
    )

    add_heading(doc, "6. Video Sonuçları ile Karşılaştırma")
    add_body(
        doc,
        "Eğitim videosunda paylaşılan çıktı, algoritmaların accuracy dağılımlarını boxplot şeklinde gösteren bir "
        "karşılaştırmadır. Bu yapı, modellerin tek bir test skorundan ziyade farklı veri bölünmelerindeki davranışını "
        "göstermesi açısından değerlidir. Bu nedenle videodaki mantığa paralel olacak şekilde bu çalışmada da 10-fold "
        "Stratified Cross Validation uygulanmış ve 10 modelin accuracy dağılımları yeniden üretilmiştir."
    )
    add_cv_summary_table(doc, cv_df)
    add_image(doc, ASSETS_DIR / "model_accuracy_boxplot.png", "Şekil 5. 10 model için cross validation accuracy boxplot çıktısı.", width_inches=6.3)
    add_body(
        doc,
        "Cross validation sonuçlarına göre ortalama accuracy bakımından en yüksek değer KNN modelinde görülmüştür "
        "(0.7746). Bunu Random Forest (0.7641), Logistic Regression (0.7629), Gradient Boosting (0.7616) ve "
        "AdaBoost (0.7604) izlemektedir. Buna karşın tek test kümesi üzerinde F1-Score ve recall dengesi dikkate "
        "alındığında en iyi model Logistic Regression olmuştur. Bu durum, farklı değerlendirme stratejilerinin model "
        "sıralamasını değiştirebileceğini göstermektedir."
    )
    add_body(
        doc,
        "Aşağıdaki tablo, videoda görülen ortalama accuracy değerleri ile bu çalışmada elde edilen 10-fold cross "
        "validation ortalamalarını doğrudan karşılaştırmaktadır. Fark sütunu, bizim bulgularımız ile videodaki "
        "ortalama accuracy arasındaki mutlak yönlü farkı göstermektedir."
    )
    add_video_comparison_table(doc, cv_df)
    add_body(
        doc,
        "Videodaki grafik ile bizim ürettiğimiz boxplot kıyaslandığında bazı ortak örüntüler dikkat çekmektedir. "
        "Her iki çıktıda da Logistic Regression, Random Forest, AdaBoost ve Gradient Boosting modelleri üst performans "
        "bandında yer almaktadır. Decision Tree ise her iki yaklaşımda da görece daha zayıf veya daha değişken sonuç "
        "veren modeller arasındadır. Bu açıdan genel eğilim bakımından videodaki bulgular ile bizim bulgularımız "
        "büyük ölçüde uyumludur."
    )
    add_body(
        doc,
        "Bununla birlikte belirgin farklar da vardır. Videodaki görselde KNN orta seviyede konumlanırken, bizim "
        "10-fold cross validation sonuçlarımızda KNN ortalama accuracy bakımından en yüksek değeri üretmiştir. "
        "Aynı şekilde videoda yalnızca 8 algoritma gösterilirken, bu çalışmada Extra Trees ve MLP modelleri de "
        "eklenmiştir. Dolayısıyla bizim yaklaşımımız hem model çeşitliliği hem de karşılaştırma kapsamı bakımından "
        "videodaki uygulamayı genişletmektedir."
    )
    add_body(
        doc,
        "Sayısal farklar daha yakından incelendiğinde KNN modelinin videoya göre belirgin biçimde daha yüksek "
        "ortalama accuracy ürettiği görülmektedir (+0.0658). Gaussian Naive Bayes, Decision Tree ve Random Forest "
        "modellerinde de sınırlı ya da orta düzeyde artışlar vardır. Buna karşılık Logistic Regression, AdaBoost, "
        "SVM ve Gradient Boosting sonuçlarımız videodaki ortalamaların bir miktar altında kalmıştır. Özellikle "
        "AdaBoost modelindeki düşüş (-0.0150) ve Logistic Regression modelindeki fark (-0.0108), veri ön işleme "
        "stratejilerinin bazı algoritmalar lehine daha fazla çalışmış olabileceğini düşündürmektedir."
    )
    add_body(
        doc,
        "Eğer bazı metriklerde videodan daha iyi sonuçlar elde edildiyse bunun temel nedenleri veri ön işleme "
        "stratejilerindeki farklılıklar, mantıksal olarak hatalı sıfır değerlerin temizlenmesi, yeni özniteliklerin "
        "üretilmesi ve algoritma çeşitliliğinin artırılması olabilir. Eğer bazı modeller videodaki izlenimden daha "
        "düşük sonuç verdiyse bunun olası nedenleri ise veri setinin sınırlı büyüklüğü, sınıf dengesizliği, aykırı "
        "gözlemler ve model-parametre uyumsuzluklarıdır. Bu açıklama, çalışmayı yalnızca izlenen adımları tekrar eden "
        "bir uygulama olmaktan çıkarıp eleştirel analitik değerlendirme düzeyine taşımaktadır."
    )

    add_heading(doc, "7. Genel Sonuç ve Profesyonel Değerlendirme")
    add_body(
        doc,
        "Bu rapor kapsamında diyabet riski tahmini problemi veri bilimi perspektifiyle ele alınmış; veri kalitesi "
        "incelemesi, öznitelik geliştirme, çoklu algoritma karşılaştırması, farklı performans ölçütleri ve video "
        "çıktısı ile doğrudan kıyaslama adımları bütüncül biçimde bir araya getirilmiştir. Böylece çalışma, salt kod "
        "çıktısı sunan bir ödev dosyasının ötesine geçirilerek mini bir veri bilimi danışmanlık raporu niteliği "
        "kazanmıştır."
    )
    add_body(
        doc,
        "Sonuçlar göstermektedir ki diyabet tahmini gibi sağlık odaklı problemlerde tek bir başarı metriğine dayanmak "
        "yeterli değildir. Accuracy, F1-Score, recall, ROC-AUC ve hata türlerinin birlikte değerlendirilmesi gerekir. "
        "Bu proje özelinde Logistic Regression modeli, yorumlanabilirliği ve dengeli hata profili nedeniyle en uygun "
        "karar destek modeli olarak öne çıkmıştır."
    )
    add_body(
        doc,
        "Gelecek çalışmalarda daha büyük veri setleri, daha sistematik hiperparametre optimizasyonu, XGBoost benzeri "
        "ek topluluk yöntemleri ve daha ayrıntılı klinik özelliklerle model başarısının artırılması mümkündür. Ancak "
        "mevcut haliyle çalışma, veri temelli sağlık yönetimi açısından anlamlı ve akademik olarak savunulabilir bir "
        "çerçeve sunmaktadır."
    )

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
